from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


def set_east_asia_font(run, name: str) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def add_inline(paragraph, text: str) -> None:
    parts = re.split(r"(\*\*.+?\*\*|`.+?`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_east_asia_font(run, "Consolas")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(55, 65, 81)
        else:
            run = paragraph.add_run(part)
        if not (part.startswith("`") and part.endswith("`")):
            set_east_asia_font(run, "Microsoft YaHei")


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    run._r.extend([field_begin, instruction, field_end])


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.2)
    add_page_number(section.footer.paragraphs[0])

    normal = document.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.35

    heading_colors = {
        "Title": RGBColor(15, 118, 110),
        "Heading 1": RGBColor(15, 118, 110),
        "Heading 2": RGBColor(17, 94, 89),
        "Heading 3": RGBColor(31, 41, 55),
        "Heading 4": RGBColor(55, 65, 81),
    }
    for style_name, color in heading_colors.items():
        style = document.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.color.rgb = color


def parse_table(lines: list[str]) -> list[list[str]]:
    rows: list[list[str]] = []
    for line in lines:
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        rows.append(cells)
    return rows


def is_separator_row(row: list[str]) -> bool:
    return all(re.fullmatch(r":?-{3,}:?", cell) for cell in row)


def add_table(document: Document, lines: list[str]) -> None:
    rows = parse_table(lines)
    rows = [row for row in rows if not is_separator_row(row)]
    if not rows:
        return
    column_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"
    for row_index, row in enumerate(rows):
        for column_index, value in enumerate(row):
            paragraph = table.cell(row_index, column_index).paragraphs[0]
            add_inline(paragraph, value)
            if row_index == 0:
                for run in paragraph.runs:
                    run.bold = True
                shading = OxmlElement("w:shd")
                shading.set(qn("w:fill"), "D9F0EC")
                table.cell(row_index, column_index)._tc.get_or_add_tcPr().append(shading)


def render_markdown(markdown_path: Path, output_path: Path) -> None:
    document = Document()
    configure_document(document)
    document.core_properties.title = "工程项目经营管理系统详细设计"
    document.core_properties.subject = "阶段 0～10 总体设计与自动执行边界"
    document.core_properties.author = "TZM / Codex"

    lines = markdown_path.read_text(encoding="utf-8").splitlines()
    index = 0
    first_heading = True
    in_code_block = False
    while index < len(lines):
        line = lines[index].rstrip()
        if line.startswith("```"):
            in_code_block = not in_code_block
            index += 1
            continue
        if in_code_block:
            paragraph = document.add_paragraph()
            run = paragraph.add_run(line)
            set_east_asia_font(run, "Consolas")
            run.font.size = Pt(8.5)
            index += 1
            continue
        if line.startswith("|"):
            table_lines: list[str] = []
            while index < len(lines) and lines[index].lstrip().startswith("|"):
                table_lines.append(lines[index])
                index += 1
            add_table(document, table_lines)
            continue
        heading = re.match(r"^(#{1,4})\s+(.+)$", line)
        if heading:
            level = len(heading.group(1))
            text = heading.group(2).strip()
            if level == 1:
                paragraph = document.add_paragraph(style="Title")
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_inline(paragraph, text)
                first_heading = False
            else:
                if level == 2 and not first_heading:
                    document.add_section(WD_SECTION.NEW_PAGE)
                paragraph = document.add_heading(level=min(level - 1, 3))
                add_inline(paragraph, text)
            index += 1
            continue
        bullet = re.match(r"^\s*-\s+(.+)$", line)
        if bullet:
            paragraph = document.add_paragraph(style="List Bullet")
            add_inline(paragraph, bullet.group(1))
            index += 1
            continue
        numbered = re.match(r"^\s*\d+\.\s+(.+)$", line)
        if numbered:
            paragraph = document.add_paragraph(style="List Number")
            add_inline(paragraph, numbered.group(1))
            index += 1
            continue
        if line.startswith("> "):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Cm(0.8)
            run = paragraph.add_run(line[2:])
            run.italic = True
            set_east_asia_font(run, "Microsoft YaHei")
            index += 1
            continue
        if line:
            paragraph = document.add_paragraph()
            add_inline(paragraph, line)
        index += 1

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


if __name__ == "__main__":
    if len(sys.argv) != 3:
        raise SystemExit("Usage: generate-design-docx.py <input.md> <output.docx>")
    render_markdown(Path(sys.argv[1]), Path(sys.argv[2]))
